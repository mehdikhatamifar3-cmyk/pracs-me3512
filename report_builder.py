from __future__ import annotations

from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# standard_business_brief preset with a restrained JCU brand override.
JCU_BLUE = "0B4F8A"
JCU_DARK_BLUE = "073B68"
JCU_GOLD = "F2A900"
INK = "172033"
MUTED = "64748B"
LIGHT_BLUE = "EAF3F9"
LIGHT_GREY = "F4F6F8"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def _set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    colour: str | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if colour is not None:
        run.font.color.rgb = RGBColor.from_string(colour)


def _shade_cell(cell, fill: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = cell_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        cell_properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in CELL_MARGIN_DXA.items():
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError("Table column widths must sum to 9360 DXA.")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_properties = table._tbl.tblPr

    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    table_width.set(qn("w:type"), "dxa")

    table_indent = table_properties.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_properties.append(table_indent)
    table_indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    table_indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440.0)
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _repeat_header(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    row_properties.append(repeat)


def _format_value(value: object) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, float):
        magnitude = abs(value)
        if magnitude and magnitude < 0.001:
            return f"{value:.3e}"
        return f"{value:.3f}".rstrip("0").rstrip(".")
    return str(value)


def _add_dataframe_table(
    document: Document,
    data: pd.DataFrame,
    columns: list[str],
    labels: list[str],
    widths_dxa: list[int],
    *,
    font_size: float = 8.0,
) -> None:
    if data.empty:
        paragraph = document.add_paragraph("No complete data are available for this section.")
        paragraph.style = document.styles["Normal"]
        return
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for index, label in enumerate(labels):
        cell = table.rows[0].cells[index]
        cell.text = label
        _shade_cell(cell, JCU_BLUE)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                _set_run_font(run, size=font_size, bold=True, colour=WHITE)
    _repeat_header(table.rows[0])
    for _, source_row in data.iterrows():
        cells = table.add_row().cells
        for index, column in enumerate(columns):
            cells[index].text = _format_value(source_row.get(column, ""))
            if len(table.rows) % 2 == 1:
                _shade_cell(cells[index], LIGHT_GREY)
            for paragraph in cells[index].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    _set_run_font(run, size=font_size, colour=INK)
    _set_table_geometry(table, widths_dxa)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def _add_labelled_paragraph(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    label_run = paragraph.add_run(f"{label}: ")
    _set_run_font(label_run, size=10.5, bold=True, colour=JCU_DARK_BLUE)
    value_run = paragraph.add_run(value)
    _set_run_font(value_run, size=10.5, colour=INK)


def _add_callout(document: Document, title: str, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _shade_cell(cell, LIGHT_BLUE)
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(f"{title}. ")
    _set_run_font(title_run, size=10.5, bold=True, colour=JCU_DARK_BLUE)
    text_run = paragraph.add_run(text)
    _set_run_font(text_run, size=10.5, colour=INK)
    _set_table_geometry(table, [TABLE_WIDTH_DXA])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_figure(
    document: Document,
    image_bytes: bytes,
    caption: str,
    explanation: str,
    figure_number: int,
) -> None:
    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.space_before = Pt(4)
    picture_paragraph.paragraph_format.space_after = Pt(4)
    picture_paragraph.paragraph_format.keep_with_next = True
    picture_run = picture_paragraph.add_run()
    picture_run.add_picture(BytesIO(image_bytes), width=Inches(6.10))

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(3)
    caption_paragraph.paragraph_format.keep_with_next = bool(explanation)
    label_run = caption_paragraph.add_run(f"Figure {figure_number}. ")
    _set_run_font(label_run, size=9.2, bold=True, colour=JCU_DARK_BLUE)
    caption_run = caption_paragraph.add_run(caption)
    _set_run_font(caption_run, size=9.2, colour=INK)

    if explanation:
        explanation_paragraph = document.add_paragraph(explanation)
        explanation_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        explanation_paragraph.paragraph_format.space_after = Pt(8)
        for run in explanation_paragraph.runs:
            _set_run_font(run, size=8.8, italic=True, colour=MUTED)


def _add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    _set_run_font(run, size=9, colour=MUTED)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, JCU_BLUE, 16, 8),
        "Heading 2": (13, JCU_BLUE, 12, 6),
        "Heading 3": (12, JCU_DARK_BLUE, 8, 4),
    }
    for style_name, (size, colour, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(colour)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_paragraph.add_run("ME3512 | ThermalLab")
    _set_run_font(run, size=9, bold=True, colour=MUTED)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_paragraph.add_run("James Cook University | Page ")
    _set_run_font(run, size=9, colour=MUTED)
    _add_page_field(footer_paragraph)


def _add_title_page(
    document: Document,
    practical_title: str,
    student_details: dict[str, str],
    logo_path: str | Path | None,
) -> None:
    if logo_path and Path(logo_path).exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(18)
        paragraph.add_run().add_picture(str(logo_path), width=Inches(1.35))

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker_run = kicker.add_run("ME3512 - HEAT AND MASS TRANSFER")
    _set_run_font(kicker_run, size=11, bold=True, colour=JCU_GOLD)
    kicker.paragraph_format.space_after = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(7)
    title_run = title.add_run(practical_title)
    _set_run_font(title_run, size=27, bold=True, colour=JCU_DARK_BLUE)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    subtitle_run = subtitle.add_run("Practical report and analysis record")
    _set_run_font(subtitle_run, size=12.5, bold=True, colour=JCU_BLUE)

    metadata = document.add_table(rows=0, cols=2)
    metadata.style = "Table Grid"
    metadata_rows = [
        ("Student", student_details.get("name", "Not entered")),
        ("JCU ID", student_details.get("student_id", "Not entered")),
        ("Group / bench", student_details.get("group", "Not entered")),
        ("Laboratory date", student_details.get("lab_date", "Not entered")),
        ("Pathway", student_details.get("pathway", "Not entered")),
        ("Report generated", datetime.now().strftime("%d %B %Y, %H:%M")),
    ]
    for label, value in metadata_rows:
        cells = metadata.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)
        _shade_cell(cells[0], LIGHT_BLUE)
        for run in cells[0].paragraphs[0].runs:
            _set_run_font(run, size=10, bold=True, colour=JCU_DARK_BLUE)
        for run in cells[1].paragraphs[0].runs:
            _set_run_font(run, size=10, colour=INK)
    _set_table_geometry(metadata, [2700, 6660])
    document.add_paragraph().paragraph_format.space_after = Pt(6)
    _add_callout(
        document,
        "Student responsibility",
        "This generated report is an analysis scaffold. Check every value, complete the discussion in your own words, and follow the LearnJCU submission instructions.",
    )
    document.add_page_break()


def _add_raw_data(document: Document, practical_code: str, raw_data: pd.DataFrame) -> None:
    _add_heading(document, "2. Raw measurements", 1)
    if practical_code == "conduction":
        _add_heading(document, "2.1 Operating conditions and T1-T4", 2)
        columns = ["Material", "Trial", "Voltage_V", "Current_A", "Water_flow_L_min", "T1_C", "T2_C", "T3_C", "T4_C"]
        labels = ["Material", "Trial", "V (V)", "I (A)", "Flow", "T1", "T2", "T3", "T4"]
        _add_dataframe_table(document, raw_data, columns, labels, [1200, 600, 800, 800, 1000, 1240, 1240, 1240, 1240], font_size=7.5)
        _add_heading(document, "2.2 T5-T8", 2)
        columns = ["Material", "Trial", "T5_C", "T6_C", "T7_C", "T8_C"]
        labels = ["Material", "Trial", "T5 (C)", "T6 (C)", "T7 (C)", "T8 (C)"]
        _add_dataframe_table(document, raw_data, columns, labels, [1900, 900, 1640, 1640, 1640, 1640], font_size=8.5)
    else:
        _add_heading(document, "2.1 Four operating cases", 2)
        first_columns = ["Case", "Fan", "Shield", "Air_velocity_m_s", "T6_air_C", "T10_wall_C"]
        first_labels = ["Case", "Fan", "Shield", "Air speed", "T6 air", "T10 wall"]
        _add_dataframe_table(document, raw_data, first_columns, first_labels, [2700, 700, 1700, 1200, 1530, 1530], font_size=7.7)
        second_columns = ["Case", "T7_polished_C", "T8_small_black_C", "T9_large_black_C"]
        second_labels = ["Case", "T7 polished", "T8 small black", "T9 large black"]
        _add_dataframe_table(document, raw_data, second_columns, second_labels, [3600, 1920, 1920, 1920], font_size=8.2)


def _add_analysis_data(document: Document, practical_code: str, analysed_data: pd.DataFrame) -> None:
    _add_heading(document, "3. Processed results", 1)
    if practical_code == "conduction":
        columns = [
            "Material",
            "Trial",
            "Assumed_conduction_heat_W",
            "Thermal_conductivity_W_mK",
            "Hot_contact_Rpp_m2K_W",
            "Cold_contact_Rpp_m2K_W",
            "Contact_share_pct",
            "Quality_flags",
        ]
        labels = ["Material", "Trial", "Q (W)", "k", "Hot R''", "Cold R''", "Contact %", "Quality check"]
        _add_dataframe_table(document, analysed_data, columns, labels, [1150, 600, 850, 900, 1050, 1050, 900, 2860], font_size=7.1)
    else:
        columns = ["Case", "T7_error_K", "T8_error_K", "T9_error_K", "Maximum_abs_error_K"]
        labels = ["Case", "T7 error", "T8 error", "T9 error", "Maximum |error|"]
        _add_dataframe_table(document, analysed_data, columns, labels, [3600, 1440, 1440, 1440, 1440], font_size=8.2)


def build_practical_report(
    *,
    practical_code: str,
    practical_title: str,
    student_details: dict[str, str],
    raw_data: pd.DataFrame,
    analysed_data: pd.DataFrame,
    aim: str,
    equations: Iterable[tuple[str, str]],
    parameter_definitions: Iterable[tuple[str, str]],
    assumptions: Iterable[tuple[str, str]],
    sample_calculation: Iterable[str],
    evidence: Iterable[str],
    discussion_notes: Iterable[tuple[str, str]],
    figures: Iterable[tuple[str, bytes, str]] = (),
    logo_path: str | Path | None = None,
) -> bytes:
    document = Document()
    _configure_document(document)
    document.core_properties.title = f"ME3512 Practical Report - {practical_title}"
    document.core_properties.subject = "ME3512 Heat and Mass Transfer practical report"
    document.core_properties.author = "James Cook University"
    _add_title_page(document, practical_title, student_details, logo_path)

    _add_heading(document, "1. Aim and governing model", 1)
    document.add_paragraph(aim)
    _add_heading(document, "1.1 Governing equations", 2)
    for equation, meaning in equations:
        _add_labelled_paragraph(document, equation, meaning)
    _add_heading(document, "1.2 Parameter definitions", 2)
    definitions = pd.DataFrame(parameter_definitions, columns=["Symbol", "Definition"])
    _add_dataframe_table(document, definitions, ["Symbol", "Definition"], ["Symbol", "Definition and unit"], [1700, 7660], font_size=9.2)
    _add_heading(document, "1.3 Analysis assumptions", 2)
    for label, value in assumptions:
        _add_labelled_paragraph(document, label, value)

    _add_raw_data(document, practical_code, raw_data)
    _add_analysis_data(document, practical_code, analysed_data)

    _add_heading(document, "4. Key graphs", 1)
    figure_items = list(figures)
    if figure_items:
        for figure_number, (caption, image_bytes, explanation) in enumerate(figure_items, start=1):
            _add_figure(document, image_bytes, caption, explanation, figure_number)
    else:
        document.add_paragraph("No complete dataset was available for generated graphs.")

    _add_heading(document, "5. Sample calculation", 1)
    if sample_calculation:
        for index, calculation in enumerate(sample_calculation, start=1):
            _add_labelled_paragraph(document, f"Step {index}", calculation)
    else:
        document.add_paragraph("No complete operating point is available for a sample calculation.")

    _add_heading(document, "6. Evidence and discussion", 1)
    for index, statement in enumerate(evidence, start=1):
        _add_labelled_paragraph(document, f"Evidence {index}", statement)
    for index, (prompt, response) in enumerate(discussion_notes, start=1):
        _add_heading(document, f"6.{index} {prompt}", 2)
        document.add_paragraph(response or "Not completed in the app.")

    _add_heading(document, "7. Conclusion and limitations", 1)
    _add_callout(
        document,
        "Before submission",
        "Summarise the main physical finding, compare the result with an appropriate reference, and distinguish random uncertainty from modelling limitations.",
    )
    for label in ("Conclusion", "Limitations and recommended improvements"):
        _add_heading(document, label, 2)
        paragraph = document.add_paragraph("Complete this section in your own words before submission.")
        for run in paragraph.runs:
            _set_run_font(run, size=10.5, italic=True, colour=MUTED)
        for _ in range(1):
            document.add_paragraph("________________________________________________________________________________")

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()
